"""
Servicio de Ingestión de Documentos.
Carga, divide y almacena documentos financieros en ChromaDB
para ser consultados mediante RAG.
"""

import os
import logging
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain.schema import Document

from app.core.config import get_settings
from app.services.guardrails import classify_financial_domain

logger = logging.getLogger(__name__)


class DocumentService:
    """
    Servicio para ingestión y gestión de documentos financieros.

    Responsabilidades:
    - Cargar documentos (PDF, TXT, DOCX)
    - Dividir en chunks optimizados para RAG
    - Almacenar embeddings en ChromaDB
    - Proveer retriever para consultas
    """

    def __init__(self):
        self.settings = get_settings()
        self._embeddings = None
        self._vector_store = None
        self._text_splitter = None

    def _create_embeddings(self) -> Any:
        """
        Factory method para crear el provider de embeddings apropiado.

        IMPORTANTE: Las importaciones son LAZY (dentro de cada rama) para evitar
        cargar librerías pesadas (torch, sentence-transformers ~400MB) cuando se
        usa un provider basado en API (google/mistral).
        """
        provider = self.settings.embedding_provider.lower()

        if provider == "google":
            if not self.settings.google_api_key:
                raise ValueError(
                    "GOOGLE_API_KEY es requerido cuando EMBEDDING_PROVIDER=google. "
                    "Obtén una gratis en https://ai.google.dev/"
                )
            from langchain_google_genai import GoogleGenerativeAIEmbeddings
            logger.info(
                f"Usando Google Generative AI Embeddings (API - GRATIS): {self.settings.google_embedding_model}"
            )
            return GoogleGenerativeAIEmbeddings(
                google_api_key=self.settings.google_api_key,
                model=self.settings.google_embedding_model,
            )

        if provider == "mistral":
            if not self.settings.mistral_api_key:
                raise ValueError(
                    "MISTRAL_API_KEY es requerido cuando EMBEDDING_PROVIDER=mistral. "
                    "Obtén una gratis en https://console.mistral.ai/"
                )
            from langchain_mistralai import MistralAIEmbeddings
            logger.info(
                f"Usando Mistral AI Embeddings (API - tier gratuito): {self.settings.mistral_embedding_model}"
            )
            return MistralAIEmbeddings(
                api_key=self.settings.mistral_api_key,
                model=self.settings.mistral_embedding_model,
            )

        # Default: HuggingFace (local - solo se importa si realmente se usa)
        from langchain_huggingface import HuggingFaceEmbeddings
        logger.info(
            f"Cargando modelo de embeddings local: {self.settings.embedding_model}"
        )
        return HuggingFaceEmbeddings(
            model_name=self.settings.embedding_model,
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )

    @property
    def embeddings(self):
        """Lazy loading de embeddings para optimizar startup."""
        if self._embeddings is None:
            self._embeddings = self._create_embeddings()
            logger.info("Embeddings inicializados exitosamente")
        return self._embeddings

    @property
    def vector_store(self) -> Chroma:
        """Lazy loading del vector store con persistencia."""
        if self._vector_store is None:
            persist_dir = self.settings.chroma_persist_dir
            os.makedirs(persist_dir, exist_ok=True)

            self._vector_store = Chroma(
                collection_name=self.settings.chroma_collection_name,
                embedding_function=self.embeddings,
                persist_directory=persist_dir,
            )
            logger.info(
                f"ChromaDB inicializado en: {persist_dir} "
                f"(colección: {self.settings.chroma_collection_name})"
            )
        return self._vector_store

    @property
    def text_splitter(self) -> RecursiveCharacterTextSplitter:
        """Text splitter configurado para documentos financieros."""
        if self._text_splitter is None:
            self._text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.settings.chunk_size,
                chunk_overlap=self.settings.chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""],
            )
        return self._text_splitter

    def _load_docx_document(self, file_path: str) -> list[Document]:
        """Carga un DOCX usando python-docx para evitar dependencias opcionales."""
        path = Path(file_path)
        docx_file = DocxDocument(str(file_path))
        content_parts: list[str] = []

        for paragraph in docx_file.paragraphs:
            text = paragraph.text.strip()
            if text:
                content_parts.append(text)

        for table in docx_file.tables:
            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_values:
                    content_parts.append(" | ".join(row_values))

        content = "\n\n".join(content_parts).strip()
        if not content:
            raise ValueError(
                f"No se encontró texto extraíble en el documento DOCX: {path.name}"
            )

        return [
            Document(
                page_content=content,
                metadata={"source": str(path), "source_file": path.name},
            )
        ]

    def _load_document(self, file_path: str) -> list[Document]:
        """
        Carga un documento según su extensión.

        Soporta: PDF, TXT, DOCX
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        supported_extensions = (".pdf", ".txt", ".docx")

        if extension not in supported_extensions:
            raise ValueError(
                f"Formato no soportado: {extension}. "
                f"Formatos válidos: {list(supported_extensions)}"
            )

        logger.info(f"Cargando documento: {path.name} ({extension})")

        if extension == ".docx":
            return self._load_docx_document(str(file_path))

        loaders = {
            ".pdf": PyPDFLoader,
            ".txt": TextLoader,
        }
        loader_class = loaders[extension]

        if extension == ".txt":
            loader = loader_class(str(file_path), encoding="utf-8")
        else:
            loader = loader_class(str(file_path))

        return loader.load()

    def ingest_document(self, file_path: str) -> dict:
        """
        Pipeline completo de ingestión: carga → split → embed → store.

        Args:
            file_path: Ruta al documento a procesar.

        Returns:
            Dict con metadata del procesamiento.
        """
        # 1. Cargar documento
        documents = self._load_document(file_path)
        logger.info(f"Documento cargado: {len(documents)} páginas/secciones")

        # 2. Dividir en chunks
        chunks = self.text_splitter.split_documents(documents)
        logger.info(f"Documento dividido en {len(chunks)} chunks")

        # 3. Filtrar contenido fuera del dominio financiero antes de indexar.
        financial_chunks: list[Document] = []
        rejected_chunks = 0
        suspicious_chunks = 0
        for chunk in chunks:
            decision = classify_financial_domain(chunk.page_content)
            if decision.allowed:
                chunk.metadata.update({
                    "domain": "finance",
                    "domain_score": decision.score,
                    "matched_terms": ", ".join(decision.matched_terms),
                    "prompt_injection_detected": decision.prompt_injection_detected,
                })
                if decision.prompt_injection_detected:
                    suspicious_chunks += 1
                    logger.warning(
                        "Posible prompt injection detectado en chunk de '%s'",
                        Path(file_path).name,
                    )
                financial_chunks.append(chunk)
            else:
                rejected_chunks += 1

        if not financial_chunks:
            raise ValueError(
                "El documento no parece pertenecer al dominio financiero. "
                "No se indexó contenido."
            )

        if rejected_chunks:
            logger.info(
                "Se descartaron %s chunks no financieros de '%s'",
                rejected_chunks,
                Path(file_path).name,
            )

        # 4. Enriquecer metadata
        filename = Path(file_path).name
        for i, chunk in enumerate(financial_chunks):
            chunk.metadata.update({
                "source_file": filename,
                "chunk_index": i,
                "total_chunks": len(financial_chunks),
            })

        # 5. Almacenar en vector store
        self.vector_store.add_documents(financial_chunks)
        logger.info(
            f"✓ {len(financial_chunks)} chunks almacenados en ChromaDB "
            f"para '{filename}'"
        )

        return {
            "filename": filename,
            "chunks_created": len(financial_chunks),
            "chunks_rejected": rejected_chunks,
            "suspicious_chunks": suspicious_chunks,
            "doc_type": Path(file_path).suffix.lstrip("."),
        }

    def ingest_directory(self, directory_path: str) -> list[dict]:
        """
        Ingesta todos los documentos soportados de un directorio.

        Args:
            directory_path: Ruta al directorio con documentos.

        Returns:
            Lista de resultados de ingestión por archivo.
        """
        results = []
        supported_extensions = {".pdf", ".txt", ".docx"}
        directory = Path(directory_path)

        if not directory.exists():
            raise FileNotFoundError(f"Directorio no encontrado: {directory_path}")

        files = [
            f for f in directory.iterdir()
            if f.is_file() and f.suffix.lower() in supported_extensions
        ]

        if not files:
            logger.warning(f"No se encontraron documentos en: {directory_path}")
            return results

        logger.info(f"Procesando {len(files)} documentos de: {directory_path}")

        for file_path in sorted(files):
            try:
                result = self.ingest_document(str(file_path))
                results.append(result)
            except Exception as e:
                logger.error(f"Error procesando {file_path.name}: {e}")
                results.append({
                    "filename": file_path.name,
                    "chunks_created": 0,
                    "doc_type": file_path.suffix.lstrip("."),
                    "error": str(e),
                })

        return results

    def get_retriever(self, k: int | None = None):
        """
        Retorna un retriever configurado para búsqueda por similitud.

        Args:
            k: Número de documentos a recuperar (default desde config).
        """
        search_k = k or self.settings.max_retrieved_docs
        return self.vector_store.as_retriever(
            search_type="similarity",
            search_kwargs={"k": search_k},
        )

    def get_collection_stats(self) -> dict:
        """Retorna estadísticas de la colección."""
        collection = self.vector_store._collection
        return {
            "total_documents": collection.count(),
            "collection_name": self.settings.chroma_collection_name,
        }

    def clear_collection(self) -> None:
        """Limpia toda la colección (útil para testing/reset)."""
        collection = self.vector_store._collection
        # Get all IDs and delete them
        all_ids = collection.get()["ids"]
        if all_ids:
            collection.delete(ids=all_ids)
        logger.info("Colección limpiada exitosamente")


# Singleton del servicio
_document_service: DocumentService | None = None


def get_document_service() -> DocumentService:
    """Factory con patrón singleton."""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service
